from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import os

def revise_document(input_path, output_path):
    doc = Document(input_path)
    
    for p in doc.paragraphs:
        original_text = p.text.strip()
        
        # Revisi 1: Tahap Analisis & Desain
        if original_text.startswith("Pada tahap ini, disusun instrumen rubrik penilaian khas Deep Learning"):
            p.clear()
            p.add_run("Pada tahap ini, disusun instrumen rubrik penilaian khas Deep Learning (Mindful, Meaningful, Joyful) bersama para ahli Teknologi Pendidikan. Rubrik ini kemudian ditransformasikan menjadi daftar kata kunci (Key Descriptors) tekstual dan parameter teknis sensor. Kemudian dikembangkan mesin NLP (Natural Language Processing) ")
            
            # Highlight hijau
            r_hl = p.add_run("dengan pendekatan Rule-Based String Matching yang disesuaikan dengan terminologi ")
            r_hl.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            
            p.add_run("pendidikan Indonesia untuk memindai dokumen RPP/Modul Ajar guru. Algoritma melakukan ekstraksi Kata Kerja Operasional (KKO) untuk mengaudit level kognitif (menargetkan PISA level 4-6) dan mencocokkan teks dengan Key Descriptors khas Deep Learning yang disusun.")
            
        # Revisi 2: Tahap Pengembangan Produk
        elif original_text.startswith("Fitur 2 berbentuk perangkat IoT di ruang kelas menggunakan unit pemroses cerdas"):
            p.clear()
            p.add_run("Fitur 2 berbentuk perangkat IoT di ruang kelas menggunakan unit pemroses cerdas (Edge Computing) mengonfigurasi perangkat keras IoT (NVIDIA Jetson, IP Camera, mikrofon array) di Sekolah Dasar mitra untuk merekam dan mengekstrak parameter pergerakan visual serta audio kelas. Data video dan audio diproses secara paralel melalui ")
            
            r_hl1 = p.add_run("model kecerdasan buatan (Artificial Intelligence) terintegrasi, ")
            r_hl1.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            
            p.add_run("yaitu (a) Computer Vision ")
            
            r_hl2 = p.add_run("(menggunakan arsitektur YOLOv8, MediaPipe, dan DeepFace) ")
            r_hl2.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            
            p.add_run("untuk ")
            
            r_hl3 = p.add_run("pelacakan postur tubuh, deteksi pergerakan manusia, dan pengenalan ekspresi wajah murid. ")
            r_hl3.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            
            p.add_run("(b) Audio Analysis: dengan menerapkan Speaker Diarization ")
            
            r_hl4 = p.add_run("(berbasis Pyannote Audio) ")
            r_hl4.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            
            p.add_run("untuk memisahkan suara guru vs murid guna menghitung Talk-Time Ratio secara objektif, serta Acoustic Event Detection untuk menangkap frekuensi tawa dan tepuk tangan murid (Joyful climate).")
            
        # Revisi 3: Analisis Mekanisme Integrasi... (YOLOv8 & MediaPipe)
        elif original_text.startswith("YOLOv8 & MediaPipe (Mindful & Meaningful):"):
            p.clear()
            p.add_run("YOLOv8 & MediaPipe (Mindful & Meaningful): Digunakan untuk deteksi manusia, analisis postur kerangka tubuh (33 keypoints landmark), serta pelacakan formasi duduk ")
            
            r_hl = p.add_run("(berdasarkan ekstraksi koordinat spasial bounding box) ")
            r_hl.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            
            p.add_run("untuk mengidentifikasi keberadaan aktivitas diskusi kelompok aktif.")
            
        # Revisi 4: Analisis Mekanisme Integrasi... (DeepFace)
        elif original_text.startswith("DeepFace & Gaze Tracking (Joyful & Mindful):"):
            p.clear()
            
            r_hl1 = p.add_run("DeepFace & Face Landmark ")
            r_hl1.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            
            p.add_run("(Joyful & Mindful): Bertanggung jawab mengklasifikasikan ekspresi emosi wajah murid secara instan (mendeteksi rasa senang/antusias) sekaligus mendeteksi arah pandangan mata murid ")
            
            r_hl2 = p.add_run("(gaze tracking) ")
            r_hl2.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            
            p.add_run("ke papan tulis atau guru sebagai metrik fokus perhatian.")

    doc.save(output_path)
    print(f"File berhasil disimpan ke {output_path}")

if __name__ == "__main__":
    input_file = r"F:\educlasify\ICETT2026_Hanif_Deep Dive Learning DSS (1).docx"
    output_file = r"F:\educlasify\ICETT2026_Hanif_Deep Dive Learning DSS_REVISI_TEGAR.docx"
    revise_document(input_file, output_file)
