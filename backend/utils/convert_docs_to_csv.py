"""
Utility script to convert uploaded PDF/DOCX documents to CSV format
Usage: python -m backend.utils.convert_docs_to_csv
"""
import re
import sys
import pandas as pd
from pathlib import Path
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from backend.database import DATABASE_URL
from backend.models.db_models import DatasetDocument


def extract_text(file_path: str) -> str:
    path = file_path.lower()
    try:
        if path.endswith(".pdf"):
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text = " ".join(p.extract_text() or "" for p in pdf.pages)
        elif path.endswith(".docx"):
            from docx import Document
            doc = Document(file_path)
            texts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    texts.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            texts.append(cell.text.strip())
            text = " ".join(texts)
        else:
            return ""
        return re.sub(r"\s+", " ", text).strip()
    except Exception:
        return ""


def convert_documents_to_csv(output_file: str = "dataset_documents.csv"):
    """
    Extract text from all uploaded PDF/DOCX files and save to CSV
    """
    engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False} if DATABASE_URL.startswith("sqlite") else {})
    Session = sessionmaker(bind=engine)
    db = Session()
    
    try:
        # Get all non-CSV documents
        docs = db.query(DatasetDocument).all()
        docs = [d for d in docs if not d.file_path.lower().endswith('.csv')]
        
        if not docs:
            print("No PDF/DOCX documents found in database")
            return
        
        print(f"Found {len(docs)} documents to convert")
        
        data = []
        for i, doc in enumerate(docs, 1):
            print(f"[{i}/{len(docs)}] Processing: {doc.file_name}")
            
            if not Path(doc.file_path).exists():
                print(f"  ⚠️  File not found: {doc.file_path}")
                continue
            
            try:
                text = extract_text(doc.file_path)
                if text:
                    data.append({
                        'text': text,
                        'label': doc.label,
                        'source_file': doc.file_name
                    })
                    print(f"  ✓ Extracted {len(text)} characters")
                else:
                    print(f"  ⚠️  No text extracted")
            except Exception as e:
                print(f"  ✗ Error: {e}")
        
        if not data:
            print("\nNo data extracted")
            return
        
        # Save to CSV
        df = pd.DataFrame(data)
        df.to_csv(output_file, index=False, encoding='utf-8')
        
        print(f"\n✓ Successfully converted {len(data)} documents")
        print(f"✓ Saved to: {output_file}")
        print(f"\nLabel distribution:")
        print(df['label'].value_counts())
        print(f"\nYou can now upload this CSV file via the admin panel")
        
    finally:
        db.close()


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Convert PDF/DOCX documents to CSV dataset")
    parser.add_argument('-o', '--output', default='dataset_documents.csv', help='Output CSV file path')
    args = parser.parse_args()
    
    convert_documents_to_csv(args.output)
