import os
import shutil
from fastapi import APIRouter, Depends, File, UploadFile, HTTPException
from sqlalchemy.orm import Session
from backend.database import get_db
from backend.models.db_models import Prediction, ActivityLog, SystemConfig
from backend.schemas.schemas import PredictionOut
from backend.services.auth_service import get_current_user
from backend.models.db_models import User
from pathlib import Path
from typing import List

router = APIRouter(tags=["predict"])

UPLOAD_VIDEO = Path("backend/uploads/video")
UPLOAD_DOC = Path("backend/uploads/document")
UPLOAD_VIDEO.mkdir(parents=True, exist_ok=True)
UPLOAD_DOC.mkdir(parents=True, exist_ok=True)





@router.post("/predict")
async def predict(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    filename = file.filename.lower()
    
    # File size validation (500MB max)
    MAX_FILE_SIZE = 500 * 1024 * 1024  # 500MB in bytes
    file_content = await file.read()
    file_size = len(file_content)
    
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400, 
            detail=f"File terlalu besar. Maksimal 500MB. Ukuran file: {file_size / (1024*1024):.1f}MB"
        )
    
    # Reset file pointer
    await file.seek(0)

    # Get confidence threshold from config
    threshold_config = db.query(SystemConfig).filter(SystemConfig.key == "confidence_threshold").first()
    confidence_threshold = float(threshold_config.value) if threshold_config else 70.0

    if filename.endswith(".mp4"):
        # Video prediction
        dest = UPLOAD_VIDEO / file.filename
        with open(dest, "wb") as f:
            f.write(file_content)

        raise HTTPException(
            status_code=501, 
            detail="Fitur prediksi video dasar (CNN) telah dinonaktifkan. Silakan gunakan fitur Analisis Video 3M."
        )

    else:
        raise HTTPException(status_code=400, detail="File harus .mp4. Untuk dokumen PDF/DOCX, gunakan fitur Analisis DLI.")

    # Check if confidence is below threshold
    low_confidence = confidence < confidence_threshold

    # Save prediction
    pred = Prediction(
        user_id=current_user.id,
        file_type=file_type,
        file_path=str(dest),
        file_name=file.filename,
        label=label,
        confidence=confidence,
    )
    db.add(pred)

    log = ActivityLog(
        user_id=current_user.id,
        action="predict",
        detail=f"{file_type} | {file.filename} → {label} ({confidence}%)",
    )
    db.add(log)
    db.commit()
    db.refresh(pred)
    
    # Cleanup old files (keep only last 100 per user)
    cleanup_old_files(db, current_user.id)

    return {
        "id": pred.id,
        "label": label,
        "confidence": confidence,
        "file_type": file_type,
        "file_name": file.filename,
        "low_confidence": low_confidence,
        "threshold": confidence_threshold,
    }


@router.post("/predict/document/detailed")
def predict_document_detailed(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Perform detailed DLI (Deep Learning Index) analysis on an RPP document.
    Returns comprehensive scores, alerts, highlighted text, and recommendations.
    
    Requirements: 12.1, 12.2, 12.3, 12.4, 13.1, 16.1, 16.3
    """
    # Authorization: only teacher, lecturer, admin roles
    ALLOWED_ROLES = {"teacher", "lecturer", "admin", "user"}
    if current_user.role not in ALLOWED_ROLES:
        raise HTTPException(
            status_code=403,
            detail="Akses ditolak. Fitur ini hanya tersedia untuk guru, dosen, dan admin."
        )

    filename = file.filename.lower()

    # Validate file format
    if not (filename.endswith(".pdf") or filename.endswith(".docx")):
        raise HTTPException(
            status_code=400,
            detail="Format file tidak didukung. Harap unggah file PDF atau DOCX."
        )

    # Validate file size (500MB max)
    MAX_FILE_SIZE = 500 * 1024 * 1024
    file_content = file.file.read()
    if len(file_content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File terlalu besar. Maksimal 500MB. Ukuran file: {len(file_content) / (1024*1024):.1f}MB"
        )

    # Save file
    dest = UPLOAD_DOC / file.filename
    with open(dest, "wb") as f:
        f.write(file_content)

    # Extract text
    try:
        import re

        def extract_text_simple(fp: str) -> str:
            try:
                if fp.lower().endswith(".pdf"):
                    import pdfplumber
                    with pdfplumber.open(fp) as pdf:
                        text = " ".join(p.extract_text() or "" for p in pdf.pages)
                elif fp.lower().endswith(".docx"):
                    from docx import Document
                    doc = Document(fp)
                    texts = []
                    for p in doc.paragraphs:
                        if p.text.strip():
                            texts.append(p.text.strip())
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    texts.append(cell.text.strip())
                    text = " ".join(texts)
                else:
                    return ""
                return re.sub(r"\s+", " ", text).strip()
            except Exception as e:
                raise RuntimeError(str(e))

        text = extract_text_simple(str(dest))
    except Exception as e:
        raise HTTPException(
            status_code=422,
            detail=f"Gagal membaca dokumen. Pastikan file tidak rusak dan dapat dibuka. Detail: {str(e)}"
        )

    # Validate minimum content
    if not text or len(text.strip()) < 100:
        raise HTTPException(
            status_code=422,
            detail="Konten dokumen terlalu sedikit untuk dianalisis. Pastikan dokumen berisi teks yang cukup (minimal 100 karakter)."
        )

    # Run DLI analysis
    try:
        from backend.ai.dli.dli_analyzer import DLIAnalyzer
        from backend.ai.dli.alert_generator import AlertGenerator
        from backend.ai.dli.recommendation_engine import RecommendationEngine
        from backend.ai.dli.text_highlighter import TextHighlighter

        analyzer = DLIAnalyzer()
        dli_result = analyzer.analyze_document(text)

        keyword_matches = dli_result.get("keyword_matches", {})

        highlighter = TextHighlighter()
        highlight_result = highlighter.highlight_text(text, keyword_matches)

        alert_gen = AlertGenerator()
        alerts = alert_gen.generate_alerts(dli_result["scores"])

        rec_engine = RecommendationEngine()
        recommendations = rec_engine.generate_recommendations(dli_result["scores"], keyword_matches)

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Terjadi kesalahan saat menganalisis dokumen. Silakan coba lagi. Detail: {str(e)}"
        )

    # Build response payload
    response_data = {
        "dli_score": dli_result["dli_score"],
        "dli_category": dli_result["dli_category"],
        "scores": dli_result["scores"],
        "sub_scores": dli_result.get("sub_scores", {}),
        "alerts": alert_gen.alerts_to_dict_list(alerts),
        "highlighted_text": highlight_result["html"],
        "keywords_found": highlight_result["keywords_found"],
        "keyword_statistics": highlight_result["statistics"],
        "recommendations": rec_engine.recommendations_to_dict_list(recommendations),
    }

    # Save to database
    pred = Prediction(
        user_id=current_user.id,
        file_type="document",
        file_path=str(dest),
        file_name=file.filename,
        label=dli_result["dli_category"],
        confidence=dli_result["dli_score"],
    )
    pred.set_dli_data({
        "dli_score": dli_result["dli_score"],
        "dli_category": dli_result["dli_category"],
        "scores": dli_result["scores"],
        "sub_scores": dli_result.get("sub_scores", {}),
        "alerts": alert_gen.alerts_to_dict_list(alerts),
        "highlighted_text": highlight_result["html"],
        "keywords_found": highlight_result["keywords_found"],
        "keyword_statistics": highlight_result["statistics"],
        "recommendations": rec_engine.recommendations_to_dict_list(recommendations),
    })
    db.add(pred)

    log = ActivityLog(
        user_id=current_user.id,
        action="dli_analysis",
        detail=f"document | {file.filename} → DLI {dli_result['dli_score']:.1f}% ({dli_result['dli_category']})",
    )
    db.add(log)
    db.commit()
    db.refresh(pred)

    cleanup_old_files(db, current_user.id)

    return {"id": pred.id, **response_data}


@router.get("/predict/history", response_model=List[PredictionOut])
def history(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    return (
        db.query(Prediction)
        .filter(Prediction.user_id == current_user.id)
        .order_by(Prediction.created_at.desc())
        .all()
    )


@router.get("/predict/export")
def export_history(
    format: str = "csv",
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Export prediction history as CSV or Excel"""
    predictions = (
        db.query(Prediction)
        .filter(Prediction.user_id == current_user.id)
        .order_by(Prediction.created_at.desc())
        .all()
    )
    
    if not predictions:
        raise HTTPException(status_code=404, detail="Tidak ada riwayat prediksi")
    
    import pandas as pd
    from io import BytesIO
    
    data = [{
        'ID': p.id,
        'Jenis File': p.file_type,
        'Nama File': p.file_name,
        'Label': p.label,
        'Confidence (%)': p.confidence,
        'Tanggal': p.created_at.strftime('%Y-%m-%d %H:%M:%S') if p.created_at else ''
    } for p in predictions]
    
    df = pd.DataFrame(data)
    
    if format == "excel":
        from fastapi.responses import StreamingResponse
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Prediction History')
        output.seek(0)
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename=prediction_history.xlsx"}
        )
    else:
        from fastapi.responses import StreamingResponse
        from io import StringIO
        csv_buffer = StringIO()
        df.to_csv(csv_buffer, index=False, encoding='utf-8')
        csv_buffer.seek(0)
        return StreamingResponse(
            iter([csv_buffer.getvalue()]),
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename=prediction_history.csv"}
        )



@router.get("/predict/document/{prediction_id}/dli")
def get_dli_result(
    prediction_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Retrieve saved DLI analysis result by prediction ID."""
    pred = db.query(Prediction).filter(
        Prediction.id == prediction_id,
        Prediction.user_id == current_user.id,
    ).first()

    if not pred:
        raise HTTPException(status_code=404, detail="Prediksi tidak ditemukan.")

    if not pred.is_dli_analysis():
        raise HTTPException(status_code=400, detail="Prediksi ini bukan hasil analisis DLI.")

    return {"id": pred.id, "file_name": pred.file_name, **pred.get_full_dli_data()}


@router.get("/predict/document/{prediction_id}/export/pdf")
def export_dli_pdf(
    prediction_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Export DLI result to PDF."""
    pred = db.query(Prediction).filter(
        Prediction.id == prediction_id,
        Prediction.user_id == current_user.id,
    ).first()

    if not pred:
        raise HTTPException(status_code=404, detail="Prediksi tidak ditemukan.")

    if not pred.is_dli_analysis():
        raise HTTPException(status_code=400, detail="Bukan hasil analisis DLI.")

    data = pred.get_full_dli_data()

    try:
        from fpdf import FPDF
    except ImportError:
        raise HTTPException(
            status_code=501,
            detail="Library PDF (fpdf2) tidak tersedia di server.",
        )

    pdf = FPDF()
    pdf.add_page()
    
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Hasil Analisis Dokumen RPP (DLI)", ln=True, align="C")
    pdf.ln(5)
    
    pdf.set_font("Helvetica", "", 12)
    pdf.cell(0, 8, f"Nama File: {pred.file_name}", ln=True)
    pdf.cell(0, 8, f"Skor Keseluruhan: {data.get('dli_score', 0):.1f}%", ln=True)
    pdf.cell(0, 8, f"Kategori: {data.get('dli_category', '')}", ln=True)
    pdf.ln(5)
    
    # Aspek Utama
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 8, "Skor Aspek Utama", ln=True)
    pdf.set_font("Helvetica", "", 12)
    scores = data.get("scores", {})
    for aspect, score in scores.items():
        pdf.cell(0, 6, f"- {aspect.capitalize()}: {score:.1f}%", ln=True)
    pdf.ln(5)
    
    # Peringatan (Alerts)
    alerts = data.get("alerts", [])
    if alerts:
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 8, "Peringatan", ln=True)
        pdf.set_font("Helvetica", "", 11)
        for alert in alerts:
            # Clean up text for PDF
            msg = alert.get("message", "").replace('→', '->')
            pdf.multi_cell(0, 6, f"[!] {msg}")
            pdf.set_x(10)
        pdf.ln(5)
        
    # Rekomendasi
    recs = data.get("recommendations", [])
    if recs:
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 8, "Rekomendasi", ln=True)
        pdf.set_font("Helvetica", "", 11)
        for rec in recs:
            title = rec.get("title", "").replace('→', '->')
            desc = rec.get("description", "").replace('→', '->')
            pdf.multi_cell(0, 6, f"- {title}: {desc}")
            pdf.set_x(10)
        pdf.ln(5)

    pdf_bytes = pdf.output()
    from fastapi.responses import StreamingResponse
    return StreamingResponse(
        iter([bytes(pdf_bytes)]),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=hasil_dli_{prediction_id}.pdf"},
    )

@router.get("/user/dashboard")
def user_dashboard(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    from datetime import datetime, timedelta
    from sqlalchemy import func
    
    # Current counts
    total = db.query(Prediction).filter(Prediction.user_id == current_user.id).count()
    video = db.query(Prediction).filter(
        Prediction.user_id == current_user.id, Prediction.file_type == "video"
    ).count()
    doc = db.query(Prediction).filter(
        Prediction.user_id == current_user.id, Prediction.file_type == "document"
    ).count()
    
    # Last 7 days trend data
    today = datetime.now().date()
    trend_data = []
    for i in range(6, -1, -1):
        day = today - timedelta(days=i)
        day_start = datetime.combine(day, datetime.min.time())
        day_end = datetime.combine(day, datetime.max.time())
        count = db.query(Prediction).filter(
            Prediction.user_id == current_user.id,
            Prediction.created_at >= day_start,
            Prediction.created_at <= day_end
        ).count()
        trend_data.append(count)
    
    # Previous week comparison
    week_ago = today - timedelta(days=7)
    two_weeks_ago = today - timedelta(days=14)
    
    this_week = db.query(Prediction).filter(
        Prediction.user_id == current_user.id,
        Prediction.created_at >= datetime.combine(week_ago, datetime.min.time())
    ).count()
    
    last_week = db.query(Prediction).filter(
        Prediction.user_id == current_user.id,
        Prediction.created_at >= datetime.combine(two_weeks_ago, datetime.min.time()),
        Prediction.created_at < datetime.combine(week_ago, datetime.min.time())
    ).count()
    
    # Calculate percentage change
    if last_week > 0:
        change_percent = round(((this_week - last_week) / last_week) * 100, 1)
    else:
        change_percent = 100.0 if this_week > 0 else 0.0
    
    # Label distribution
    label_dist = db.query(
        Prediction.label,
        func.count(Prediction.id).label('count')
    ).filter(
        Prediction.user_id == current_user.id
    ).group_by(Prediction.label).all()
    
    label_distribution = {label: count for label, count in label_dist}
    
    recent = (
        db.query(Prediction)
        .filter(Prediction.user_id == current_user.id)
        .order_by(Prediction.created_at.desc())
        .limit(5)
        .all()
    )
    
    return {
        "total_predictions": total,
        "video_predictions": video,
        "document_predictions": doc,
        "trend_data": trend_data,
        "change_percent": change_percent,
        "this_week_count": this_week,
        "last_week_count": last_week,
        "label_distribution": label_distribution,
        "achievements": get_achievements(total, this_week),
        "recent": [
            {
                "id": p.id,
                "file_type": p.file_type,
                "file_name": p.file_name,
                "label": p.label,
                "confidence": p.confidence,
                "created_at": p.created_at,
            }
            for p in recent
        ],
    }


def get_achievements(total: int, this_week: int) -> list:
    """Get user achievements based on activity"""
    achievements = []
    
    # Total predictions milestones
    if total >= 100:
        achievements.append({"icon": "🏆", "title": "Century", "desc": "100+ prediksi"})
    elif total >= 50:
        achievements.append({"icon": "🥇", "title": "Gold", "desc": "50+ prediksi"})
    elif total >= 10:
        achievements.append({"icon": "🥈", "title": "Silver", "desc": "10+ prediksi"})
    elif total >= 1:
        achievements.append({"icon": "🥉", "title": "Bronze", "desc": "Prediksi pertama"})
    
    # Weekly activity
    if this_week >= 20:
        achievements.append({"icon": "🔥", "title": "On Fire", "desc": "20+ minggu ini"})
    elif this_week >= 10:
        achievements.append({"icon": "⚡", "title": "Active", "desc": "10+ minggu ini"})
    
    return achievements



@router.delete("/predict/{prediction_id}")
def delete_prediction(
    prediction_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Delete a prediction by ID"""
    prediction = db.query(Prediction).filter(
        Prediction.id == prediction_id,
        Prediction.user_id == current_user.id
    ).first()
    
    if not prediction:
        raise HTTPException(status_code=404, detail="Prediksi tidak ditemukan")
    
    # Delete file if exists
    if prediction.file_path and os.path.exists(prediction.file_path):
        try:
            os.remove(prediction.file_path)
        except:
            pass
    
    # Delete from database
    db.delete(prediction)
    
    # Log activity
    log = ActivityLog(
        user_id=current_user.id,
        action="delete_prediction",
        detail=f"Deleted: {prediction.file_name}",
    )
    db.add(log)
    db.commit()
    
    return {"message": "Prediksi berhasil dihapus"}



def cleanup_old_files(db: Session, user_id: int, keep_count: int = 100):
    """Delete old prediction files to save storage"""
    # Get all predictions for user, ordered by date
    predictions = db.query(Prediction).filter(
        Prediction.user_id == user_id
    ).order_by(Prediction.created_at.desc()).all()
    
    # If more than keep_count, delete the oldest ones
    if len(predictions) > keep_count:
        to_delete = predictions[keep_count:]
        
        for pred in to_delete:
            # Delete file if exists
            if pred.file_path and os.path.exists(pred.file_path):
                try:
                    os.remove(pred.file_path)
                except:
                    pass
            
            # Delete from database
            db.delete(pred)
        
        db.commit()


@router.get("/storage/stats")
def get_storage_stats(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Get storage usage statistics for current user"""
    predictions = db.query(Prediction).filter(
        Prediction.user_id == current_user.id
    ).all()
    
    total_size = 0
    video_size = 0
    document_size = 0
    video_count = 0
    document_count = 0
    
    for pred in predictions:
        if pred.file_path and os.path.exists(pred.file_path):
            file_size = os.path.getsize(pred.file_path)
            total_size += file_size
            
            if pred.file_type == 'video':
                video_size += file_size
                video_count += 1
            else:
                document_size += file_size
                document_count += 1
    
    return {
        "total_size_mb": round(total_size / (1024 * 1024), 2),
        "video_size_mb": round(video_size / (1024 * 1024), 2),
        "document_size_mb": round(document_size / (1024 * 1024), 2),
        "video_count": video_count,
        "document_count": document_count,
        "total_predictions": len(predictions),
    }


@router.post("/storage/cleanup")
def manual_cleanup(
    keep_count: int = 50,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Manually trigger cleanup of old files"""
    predictions_before = db.query(Prediction).filter(
        Prediction.user_id == current_user.id
    ).count()
    
    cleanup_old_files(db, current_user.id, keep_count)
    
    predictions_after = db.query(Prediction).filter(
        Prediction.user_id == current_user.id
    ).count()
    
    deleted_count = predictions_before - predictions_after
    
    return {
        "message": f"Cleanup selesai. {deleted_count} file lama dihapus.",
        "deleted_count": deleted_count,
        "remaining_count": predictions_after,
    }
