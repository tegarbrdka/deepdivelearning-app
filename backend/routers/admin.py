import asyncio
import uuid
import json
import shutil
from pathlib import Path
from fastapi import APIRouter, Depends, File, UploadFile, HTTPException, Form, WebSocket, WebSocketDisconnect
from sqlalchemy.orm import Session
from sqlalchemy import func
from typing import List, Optional
from backend.database import get_db
from backend.models.db_models import (
    User, Prediction, DatasetVideo, DatasetDocument,
    TrainingLog,
    ActivityLog, SystemConfig
)
from backend.schemas.schemas import (
    AdminStats, ActivityLogOut, SystemConfigUpdate,
    UserCreate, UserUpdate, ModelVersionOut,
    DatasetVideoOut, DatasetDocumentOut
)
from backend.services.auth_service import require_admin, hash_password

router = APIRouter(prefix="/admin", tags=["admin"])

UPLOAD_VIDEO = Path("backend/uploads/video")
UPLOAD_DOC = Path("backend/uploads/document")
UPLOAD_VIDEO.mkdir(parents=True, exist_ok=True)
UPLOAD_DOC.mkdir(parents=True, exist_ok=True)

# In-memory training job tracker
training_jobs: dict = {}


# ── Stats ─────────────────────────────────────────────────────────────────────
@router.get("/stats")
def admin_stats(db: Session = Depends(get_db), _=Depends(require_admin)):
    try:
        from datetime import datetime, timedelta
        total_preds = db.query(Prediction).count()
        total_users = db.query(User).count()
        total_video_ds = db.query(DatasetVideo).count()
        total_doc_ds = db.query(DatasetDocument).count()

        # Daily predictions last 7 days
        today = datetime.utcnow().date()
        daily = []
        for i in range(6, -1, -1):
            day = today - timedelta(days=i)
            count = db.query(Prediction).filter(
                func.date(Prediction.created_at) == day
            ).count()
            daily.append({"date": day.strftime("%d/%m"), "count": count})

        # DLI stats from Prediction table
        dli_total = db.query(Prediction).filter(Prediction.dli_score.isnot(None)).count()
        dli_avg = db.query(func.avg(Prediction.dli_score)).filter(Prediction.dli_score.isnot(None)).scalar()

        # Label distribution
        dl_count = db.query(Prediction).filter(Prediction.label == "Deep Learning").count()
        not_dl_count = db.query(Prediction).filter(Prediction.label == "Bukan Deep Learning").count()

        return {
            "total_predictions": total_preds,
            "total_users": total_users,
            "total_video_dataset": total_video_ds,
            "total_document_dataset": total_doc_ds,

            "daily_predictions": daily,
            "dli_total": dli_total,
            "dli_avg_score": round(dli_avg, 1) if dli_avg else None,
            "label_distribution": {"deep_learning": dl_count, "not_deep_learning": not_dl_count},
        }
    except Exception as e:
        print(f"Stats error: {e}")
        import traceback
        traceback.print_exc()
        return {
            "total_predictions": 0,
            "total_users": 0,
            "total_video_dataset": 0,
            "total_document_dataset": 0,

            "daily_predictions": [],
            "dli_total": 0,
            "dli_avg_score": None,
            "label_distribution": {"deep_learning": 0, "not_deep_learning": 0},
        }


# ── Dataset Video ─────────────────────────────────────────────────────────────
@router.get("/dataset/video", response_model=List[DatasetVideoOut])
def list_video_dataset(
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    _=Depends(require_admin)
):
    return db.query(DatasetVideo).order_by(DatasetVideo.created_at.desc()).offset(skip).limit(limit).all()


@router.post("/dataset/video")
async def upload_video_dataset(
    file: UploadFile = File(...),
    label: str = Form(...),
    group_name: str = Form("Default"),
    db: Session = Depends(get_db),
    admin=Depends(require_admin),
):
    if not file.filename.lower().endswith(".mp4"):
        raise HTTPException(status_code=400, detail="File harus .mp4")

    # Sanitize filename to avoid issues with spaces/special chars
    safe_name = file.filename.replace(" ", "_")
    dest = UPLOAD_VIDEO / f"ds_{uuid.uuid4().hex}_{safe_name}"

    # Stream file in chunks to support large video files
    try:
        with open(dest, "wb") as out:
            while chunk := await file.read(1024 * 1024):  # 1MB chunks
                out.write(chunk)
    except Exception as e:
        if dest.exists():
            dest.unlink()
        raise HTTPException(status_code=500, detail=f"Gagal menyimpan file: {str(e)}")

    item = DatasetVideo(
        file_path=str(dest),
        file_name=file.filename,
        label=label,
        group_name=group_name,
        uploaded_by=admin.id,
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return {"id": item.id, "file_name": item.file_name, "label": item.label, "group_name": item.group_name}


@router.delete("/dataset/video/{item_id}")
def delete_video_dataset(item_id: int, db: Session = Depends(get_db), _=Depends(require_admin)):
    item = db.query(DatasetVideo).filter(DatasetVideo.id == item_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="Not found")
    if item.file_path and Path(item.file_path).exists():
        Path(item.file_path).unlink()
    db.delete(item)
    db.commit()
    return {"message": "Deleted"}


# ── Dataset Document ──────────────────────────────────────────────────────────
@router.get("/dataset/document", response_model=List[DatasetDocumentOut])
def list_doc_dataset(
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    _=Depends(require_admin)
):
    return db.query(DatasetDocument).order_by(DatasetDocument.created_at.desc()).offset(skip).limit(limit).all()


@router.post("/dataset/document")
async def upload_doc_dataset(
    file: UploadFile = File(...),
    label: str = Form(None),
    group_name: str = Form("Default"),
    db: Session = Depends(get_db),
    admin=Depends(require_admin),
):
    fname = file.filename.lower()
    
    # Handle CSV upload
    if fname.endswith(".csv"):
        import pandas as pd
        from io import StringIO
        
        content = await file.read()
        try:
            df = pd.read_csv(StringIO(content.decode('utf-8')))
        except:
            raise HTTPException(status_code=400, detail="CSV tidak valid")
        
        if 'text' not in df.columns or 'label' not in df.columns:
            raise HTTPException(status_code=400, detail="CSV harus memiliki kolom 'text' dan 'label'")
        
        # Validate labels
        valid_labels = ["Baik", "Cukup", "Kurang"]
        invalid = df[~df['label'].isin(valid_labels)]
        if len(invalid) > 0:
            raise HTTPException(status_code=400, detail=f"Label tidak valid. Harus: {', '.join(valid_labels)}")
        
        # Save CSV to uploads
        dest = UPLOAD_DOC / f"ds_{uuid.uuid4().hex}_{file.filename}"
        with open(dest, "wb") as f:
            f.write(content)
        
        # Create single dataset entry pointing to CSV
        item = DatasetDocument(
            file_path=str(dest), 
            file_name=file.filename, 
            label=f"CSV ({len(df)} rows)", 
            group_name=group_name,
            uploaded_by=admin.id
        )
        db.add(item)
        db.commit()
        db.refresh(item)
        return {"id": item.id, "file_name": item.file_name, "label": item.label, "group_name": item.group_name, "rows": len(df)}
    
    # Handle PDF/DOCX upload (existing logic)
    if not (fname.endswith(".pdf") or fname.endswith(".docx")):
        raise HTTPException(status_code=400, detail="File harus .pdf, .docx, atau .csv")
    
    if not label:
        raise HTTPException(status_code=400, detail="Label wajib untuk file PDF/DOCX")
    
    dest = UPLOAD_DOC / f"ds_{uuid.uuid4().hex}_{file.filename}"
    with open(dest, "wb") as f:
        shutil.copyfileobj(file.file, f)
    item = DatasetDocument(file_path=str(dest), file_name=file.filename, label=label, group_name=group_name, uploaded_by=admin.id)
    db.add(item)
    db.commit()
    db.refresh(item)
    return {"id": item.id, "file_name": item.file_name, "label": item.label, "group_name": item.group_name}


@router.delete("/dataset/document/{item_id}")
def delete_doc_dataset(item_id: int, db: Session = Depends(get_db), _=Depends(require_admin)):
    item = db.query(DatasetDocument).filter(DatasetDocument.id == item_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="Not found")
    if item.file_path and Path(item.file_path).exists():
        Path(item.file_path).unlink()
    db.delete(item)
    db.commit()
    return {"message": "Deleted"}



# ── Users ─────────────────────────────────────────────────────────────────────
@router.get("/users")
def list_users(
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    _=Depends(require_admin)
):
    users = db.query(User).offset(skip).limit(limit).all()
    return [
        {"id": u.id, "username": u.username, "email": u.email, "role": u.role, "created_at": u.created_at}
        for u in users
    ]


@router.post("/users")
def create_user(payload: UserCreate, db: Session = Depends(get_db), _=Depends(require_admin)):
    if db.query(User).filter(User.username == payload.username).first():
        raise HTTPException(status_code=400, detail="Username exists")
    user = User(
        username=payload.username, email=payload.email,
        password_hash=hash_password(payload.password), role=payload.role
    )
    db.add(user)
    db.commit()
    return {"message": "User created", "id": user.id}


@router.delete("/users/{user_id}")
def delete_user(user_id: int, db: Session = Depends(get_db), _=Depends(require_admin)):
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    db.delete(user)
    db.commit()
    return {"message": "User deleted"}


@router.patch("/users/{user_id}")
def update_user(user_id: int, payload: UserUpdate, db: Session = Depends(get_db), _=Depends(require_admin)):
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    if payload.role:
        user.role = payload.role
    if payload.password:
        user.password_hash = hash_password(payload.password)
    db.commit()
    return {"message": "User updated"}


# ── Logs ──────────────────────────────────────────────────────────────────────
@router.get("/logs")
def get_logs(
    skip: int = 0, limit: int = 50,
    db: Session = Depends(get_db), _=Depends(require_admin)
):
    logs = (
        db.query(ActivityLog, User.username)
        .outerjoin(User, ActivityLog.user_id == User.id)
        .order_by(ActivityLog.created_at.desc())
        .offset(skip).limit(limit).all()
    )
    return [
        {
            "id": l.ActivityLog.id,
            "username": l.username,
            "action": l.ActivityLog.action,
            "detail": l.ActivityLog.detail,
            "created_at": l.ActivityLog.created_at,
        }
        for l in logs
    ]


# ── Export ────────────────────────────────────────────────────────────────────
@router.get("/export/dataset")
def export_dataset(db: Session = Depends(get_db), _=Depends(require_admin)):
    import pandas as pd
    from fastapi.responses import StreamingResponse
    import io

    videos = db.query(DatasetVideo).all()
    docs = db.query(DatasetDocument).all()

    data = []
    for v in videos:
        data.append({"type": "video", "file_name": v.file_name, "label": v.label, "created_at": v.created_at})
    for d in docs:
        data.append({"type": "document", "file_name": d.file_name, "label": d.label, "created_at": d.created_at})

    df = pd.DataFrame(data)
    buf = io.BytesIO()
    df.to_excel(buf, index=False)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=dataset.xlsx"}
    )


@router.get("/export/logs")
def export_logs(db: Session = Depends(get_db), _=Depends(require_admin)):
    import pandas as pd
    from fastapi.responses import StreamingResponse
    import io

    logs = (
        db.query(ActivityLog, User.username)
        .outerjoin(User, ActivityLog.user_id == User.id)
        .order_by(ActivityLog.created_at.desc())
        .all()
    )
    data = [
        {"username": l.username, "action": l.ActivityLog.action,
         "detail": l.ActivityLog.detail, "created_at": l.ActivityLog.created_at}
        for l in logs
    ]
    df = pd.DataFrame(data)
    buf = io.BytesIO()
    df.to_csv(buf, index=False)
    buf.seek(0)
    return StreamingResponse(
        buf, media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=logs.csv"}
    )


# ── Export Documents to CSV ───────────────────────────────────────────────────
@router.get("/dataset/export-csv")
def export_documents_to_csv(db: Session = Depends(get_db), _=Depends(require_admin)):
    """Export all PDF/DOCX documents to CSV format"""
    import re
    import pandas as pd
    from io import StringIO

    def extract_text(file_path: str) -> str:
        path = file_path.lower()
        try:
            if path.endswith(".pdf"):
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    text = " ".join(p.extract_text() or "" for p in pdf.pages)
            elif path.endswith(".docx"):
                from docx import Document
                doc = Document(file_path)
                texts = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        texts.append(p.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                texts.append(cell.text.strip())
                text = " ".join(texts)
            else:
                return ""
            text = re.sub(r"\s+", " ", text).strip()
            return text
        except Exception:
            return ""

    docs = db.query(DatasetDocument).all()
    docs = [d for d in docs if not d.file_path.lower().endswith('.csv')]

    if not docs:
        raise HTTPException(status_code=404, detail="No PDF/DOCX documents found")

    data = []
    for doc in docs:
        if Path(doc.file_path).exists():
            text = extract_text(doc.file_path)
            if text:
                data.append({'text': text, 'label': doc.label, 'source_file': doc.file_name})

    if not data:
        raise HTTPException(status_code=404, detail="No text could be extracted")

    df = pd.DataFrame(data)
    csv_buffer = StringIO()
    df.to_csv(csv_buffer, index=False, encoding='utf-8')

    from fastapi.responses import StreamingResponse
    csv_buffer.seek(0)
    return StreamingResponse(
        iter([csv_buffer.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=dataset_documents.csv"}
    )


# ── System Config ─────────────────────────────────────────────────────────────
@router.get("/config")
def get_config(db: Session = Depends(get_db), _=Depends(require_admin)):
    configs = db.query(SystemConfig).all()
    return {c.key: c.value for c in configs}


@router.post("/config")
def update_config(payload: SystemConfigUpdate, db: Session = Depends(get_db), _=Depends(require_admin)):
    cfg = db.query(SystemConfig).filter(SystemConfig.key == payload.key).first()
    if cfg:
        cfg.value = payload.value
    else:
        cfg = SystemConfig(key=payload.key, value=payload.value)
        db.add(cfg)
    db.commit()
    return {"message": "Config updated"}


# ── Model Cache ───────────────────────────────────────────────────────────────
@router.get("/cache/info")
def get_cache_info(_=Depends(require_admin)):
    """Get model cache information"""
    from backend.ai.model_cache import model_cache
    return model_cache.get_cache_info()


@router.post("/cache/clear")
def clear_cache(_=Depends(require_admin)):
    """Clear model cache (force reload on next prediction)"""
    from backend.ai.model_cache import model_cache
    model_cache.clear_cache()
    return {"message": "Model cache cleared successfully"}


# ── DLI Dashboard ─────────────────────────────────────────────────────────────
@router.get("/dli/dashboard")
def dli_dashboard(db: Session = Depends(get_db), _=Depends(require_admin)):
    """Aggregated DLI statistics for admin dashboard"""
    preds = db.query(Prediction).filter(Prediction.dli_score.isnot(None)).all()

    if not preds:
        return {
            "total": 0, "avg_score": 0,
            "grade_distribution": {"grade4": 0, "grade3": 0, "grade2": 0, "grade1": 0},
            "aspect_averages": {},
            "weakest_aspects": [],
            "strongest_aspects": [],
            "recent_trend": [],
            "category_counts": {},
        }

    total = len(preds)
    scores = [p.dli_score for p in preds]
    avg_score = round(sum(scores) / total, 1)

    grade_dist = {"grade4": 0, "grade3": 0, "grade2": 0, "grade1": 0}
    for s in scores:
        if s >= 70:   grade_dist["grade4"] += 1
        elif s >= 55: grade_dist["grade3"] += 1
        elif s >= 40: grade_dist["grade2"] += 1
        else:         grade_dist["grade1"] += 1

    aspect_fields = {
        "mindful": "mindful_score", "meaningful": "meaningful_score",
        "joyful": "joyful_score", "pedagogis": "pedagogis_score", "digital": "digital_score",
    }
    aspect_avgs = {}
    for aspect, field in aspect_fields.items():
        vals = [getattr(p, field) for p in preds if getattr(p, field) is not None]
        aspect_avgs[aspect] = round(sum(vals) / len(vals), 1) if vals else 0

    sorted_aspects = sorted(aspect_avgs.items(), key=lambda x: x[1])
    weakest   = [{"aspect": a, "avg": v} for a, v in sorted_aspects[:2]]
    strongest = [{"aspect": a, "avg": v} for a, v in sorted_aspects[-2:][::-1]]

    cat_counts = {}
    for p in preds:
        cat = p.dli_category or "Unknown"
        cat_counts[cat] = cat_counts.get(cat, 0) + 1

    recent = sorted(preds, key=lambda p: p.created_at or 0, reverse=True)[:10]
    trend = [
        {"file": p.file_name or f"ID:{p.id}", "score": round(p.dli_score, 1), "date": str(p.created_at)[:10]}
        for p in reversed(recent)
    ]

    return {
        "total": total, "avg_score": avg_score,
        "grade_distribution": grade_dist,
        "aspect_averages": aspect_avgs,
        "weakest_aspects": weakest,
        "strongest_aspects": strongest,
        "recent_trend": trend,
        "category_counts": cat_counts,
    }


# ── DLI History ────────────────────────────────────────────────────────────────
@router.get("/dli/history")
def dli_history(
    grade: str = None,       # "grade4" | "grade3" | "grade2" | "grade1"
    user_id: int = None,
    date_from: str = None,   # YYYY-MM-DD
    date_to: str = None,
    limit: int = 100,
    db: Session = Depends(get_db),
    _=Depends(require_admin),
):
    """All DLI analyses with filters"""
    from datetime import datetime
    q = db.query(Prediction, User.username).outerjoin(User, Prediction.user_id == User.id)\
        .filter(Prediction.dli_score.isnot(None))

    if user_id:
        q = q.filter(Prediction.user_id == user_id)
    if date_from:
        q = q.filter(Prediction.created_at >= datetime.strptime(date_from, "%Y-%m-%d"))
    if date_to:
        q = q.filter(Prediction.created_at <= datetime.strptime(date_to + " 23:59:59", "%Y-%m-%d %H:%M:%S"))
    if grade:
        ranges = {"grade4": (70, 101), "grade3": (55, 70), "grade2": (40, 55), "grade1": (0, 40)}
        if grade in ranges:
            lo, hi = ranges[grade]
            q = q.filter(Prediction.dli_score >= lo, Prediction.dli_score < hi)

    rows = q.order_by(Prediction.created_at.desc()).limit(limit).all()
    return [
        {
            "id": p.Prediction.id,
            "file_name": p.Prediction.file_name,
            "username": p.username,
            "dli_score": round(p.Prediction.dli_score, 1),
            "dli_category": p.Prediction.dli_category,
            "mindful": round(p.Prediction.mindful_score or 0, 1),
            "meaningful": round(p.Prediction.meaningful_score or 0, 1),
            "joyful": round(p.Prediction.joyful_score or 0, 1),
            "pedagogis": round(p.Prediction.pedagogis_score or 0, 1),
            "digital": round(p.Prediction.digital_score or 0, 1),
            "created_at": str(p.Prediction.created_at)[:19],
        }
        for p in rows
    ]


@router.get("/dli/history/export")
def export_dli_history(
    grade: str = None,
    user_id: int = None,
    date_from: str = None,
    date_to: str = None,
    db: Session = Depends(get_db),
    _=Depends(require_admin),
):
    """Export DLI history as Excel"""
    import pandas as pd
    from fastapi.responses import StreamingResponse
    from io import BytesIO
    from datetime import datetime

    q = db.query(Prediction, User.username).outerjoin(User, Prediction.user_id == User.id)\
        .filter(Prediction.dli_score.isnot(None))
    if user_id:
        q = q.filter(Prediction.user_id == user_id)
    if date_from:
        q = q.filter(Prediction.created_at >= datetime.strptime(date_from, "%Y-%m-%d"))
    if date_to:
        q = q.filter(Prediction.created_at <= datetime.strptime(date_to + " 23:59:59", "%Y-%m-%d %H:%M:%S"))
    if grade:
        ranges = {"grade4": (70, 101), "grade3": (55, 70), "grade2": (40, 55), "grade1": (0, 40)}
        if grade in ranges:
            lo, hi = ranges[grade]
            q = q.filter(Prediction.dli_score >= lo, Prediction.dli_score < hi)

    rows = q.order_by(Prediction.created_at.desc()).all()
    data = [{
        "ID": p.Prediction.id, "File": p.Prediction.file_name, "User": p.username,
        "DLI Score": round(p.Prediction.dli_score, 1), "Category": p.Prediction.dli_category,
        "Mindful": round(p.Prediction.mindful_score or 0, 1),
        "Meaningful": round(p.Prediction.meaningful_score or 0, 1),
        "Joyful": round(p.Prediction.joyful_score or 0, 1),
        "Pedagogis": round(p.Prediction.pedagogis_score or 0, 1),
        "Digital": round(p.Prediction.digital_score or 0, 1),
        "Date": str(p.Prediction.created_at)[:19],
    } for p in rows]

    df = pd.DataFrame(data)
    buf = BytesIO()
    df.to_excel(buf, index=False)
    buf.seek(0)
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=dli_history.xlsx"})


# ── DLI Analytics ──────────────────────────────────────────────────────────────
@router.get("/dli/analytics")
def dli_analytics(db: Session = Depends(get_db), _=Depends(require_admin)):
    """Advanced DLI analytics: heatmap + keyword effectiveness"""
    try:
        preds = db.query(Prediction).filter(Prediction.dli_score.isnot(None)).all()
        if not preds:
            return {"heatmap": {}, "keyword_effectiveness": [], "grade_counts": {}}

        grades = {"grade4": [], "grade3": [], "grade2": [], "grade1": []}
        for p in preds:
            if p.dli_score >= 70:   g = "grade4"
            elif p.dli_score >= 55: g = "grade3"
            elif p.dli_score >= 40: g = "grade2"
            else:                   g = "grade1"
            grades[g].append(p)

        aspects = ["mindful", "meaningful", "joyful", "pedagogis", "digital"]
        heatmap = {}
        for grade, ps in grades.items():
            heatmap[grade] = {}
            for a in aspects:
                field = f"{a}_score"
                vals = [getattr(p, field) for p in ps if getattr(p, field) is not None]
                heatmap[grade][a] = round(sum(vals) / len(vals), 1) if vals else 0

        keyword_freq = {}
        for p in grades["grade4"]:
            try:
                full = p.get_full_dli_data() or {}
                kw_found = full.get("keywords_found", {})
                if not isinstance(kw_found, dict):
                    continue
                for kw_list in kw_found.values():
                    if isinstance(kw_list, list):
                        for kw in kw_list:
                            if isinstance(kw, str) and kw.strip():
                                keyword_freq[kw] = keyword_freq.get(kw, 0) + 1
            except Exception:
                continue

        top_keywords = sorted(keyword_freq.items(), key=lambda x: x[1], reverse=True)[:20]

        return {
            "heatmap": heatmap,
            "keyword_effectiveness": [{"keyword": k, "count": v} for k, v in top_keywords],
            "grade_counts": {g: len(ps) for g, ps in grades.items()},
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Analytics error: {str(e)}")


# ── DLI Bulk Analysis ──────────────────────────────────────────────────────────
@router.post("/dli/bulk-analyze")
async def bulk_analyze(
    files: List[UploadFile] = File(...),
    db: Session = Depends(get_db),
    admin=Depends(require_admin),
):
    """Analyze multiple RPP documents at once"""
    import re

    def extract_text_simple(fp: str) -> tuple:
        """Returns (text, error_message)"""
        try:
            if fp.lower().endswith(".pdf"):
                import pdfplumber
                with pdfplumber.open(fp) as pdf:
                    if len(pdf.pages) == 0:
                        return "", "PDF tidak memiliki halaman"
                    text = " ".join(p.extract_text() or "" for p in pdf.pages)
                    if not text.strip():
                        return "", "PDF mungkin berupa scan/gambar. Gunakan OCR atau convert ke text-based PDF"
            elif fp.lower().endswith(".docx"):
                from docx import Document as DocxDoc
                doc = DocxDoc(fp)
                paragraphs = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        paragraphs.append(p.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                paragraphs.append(cell.text.strip())
                if not paragraphs:
                    return "", "Dokumen Word kosong atau tidak memiliki teks"
                text = " ".join(paragraphs)
            else:
                return "", "Format file tidak didukung"
            
            cleaned = re.sub(r"\s+", " ", text).strip()
            if len(cleaned) < 50:
                return "", f"Teks terlalu pendek ({len(cleaned)} karakter). Minimal 50 karakter diperlukan"
            
            return cleaned, None
        except Exception as e:
            return "", f"Error membaca file: {str(e)}"

    results = []
    for file in files:
        fname = file.filename.lower()
        if not (fname.endswith(".pdf") or fname.endswith(".docx")):
            results.append({"file": file.filename, "error": "Unsupported format"})
            continue

        dest = UPLOAD_DOC / f"bulk_{uuid.uuid4().hex[:8]}_{file.filename}"
        content = await file.read()
        with open(dest, "wb") as f:
            f.write(content)

        text, error_msg = extract_text_simple(str(dest))
        if error_msg:
            results.append({"file": file.filename, "error": error_msg})
            continue

        try:
            from backend.ai.dli.dli_analyzer import DLIAnalyzer
            from backend.ai.dli.alert_generator import AlertGenerator
            from backend.ai.dli.recommendation_engine import RecommendationEngine

            analyzer = DLIAnalyzer()
            dli_result = analyzer.analyze_document(text)
            alert_gen = AlertGenerator()
            alerts = alert_gen.generate_alerts(dli_result["scores"])
            rec_engine = RecommendationEngine()
            recommendations = rec_engine.generate_recommendations(dli_result["scores"], dli_result.get("keyword_matches", {}))

            pred = Prediction(
                user_id=admin.id, file_type="document",
                file_path=str(dest), file_name=file.filename,
                label=dli_result["dli_category"], confidence=dli_result["dli_score"],
            )
            pred.set_dli_data({
                "dli_score": dli_result["dli_score"],
                "dli_category": dli_result["dli_category"],
                "scores": dli_result["scores"],
                "sub_scores": dli_result.get("sub_scores", {}),
                "alerts": alert_gen.alerts_to_dict_list(alerts),
                "recommendations": rec_engine.recommendations_to_dict_list(recommendations),
                "highlighted_text": "", "keywords_found": {}, "keyword_statistics": {},
            })
            db.add(pred)
            db.commit()
            db.refresh(pred)

            results.append({
                "id": pred.id,
                "file": file.filename,
                "dli_score": dli_result["dli_score"],
                "dli_category": dli_result["dli_category"],
                "scores": dli_result["scores"],
            })
        except Exception as e:
            results.append({"file": file.filename, "error": str(e)})

    return {"results": results, "total": len(files), "success": sum(1 for r in results if "error" not in r)}


# ── DLI Keyword Management ────────────────────────────────────────────────────
KEYWORDS_DIR = Path("backend/ai/dli/keywords")
VALID_ASPECTS = ["mindful", "meaningful", "joyful", "pedagogis", "digital"]


@router.get("/dli/keywords")
def get_all_keywords(_=Depends(require_admin)):
    """Get all keyword dictionaries for all aspects"""
    result = {}
    for aspect in VALID_ASPECTS:
        file_path = KEYWORDS_DIR / f"{aspect}.json"
        if file_path.exists():
            with open(file_path, "r", encoding="utf-8") as f:
                result[aspect] = json.load(f)
    return result


@router.get("/dli/keywords/{aspect}")
def get_aspect_keywords(aspect: str, _=Depends(require_admin)):
    """Get keyword dictionary for a specific aspect"""
    if aspect not in VALID_ASPECTS:
        raise HTTPException(status_code=400, detail=f"Invalid aspect. Must be one of: {VALID_ASPECTS}")
    file_path = KEYWORDS_DIR / f"{aspect}.json"
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Keyword file not found")
    with open(file_path, "r", encoding="utf-8") as f:
        return json.load(f)


@router.put("/dli/keywords/{aspect}")
def update_aspect_keywords(aspect: str, payload: dict, _=Depends(require_admin)):
    """Replace entire keyword dictionary for a specific aspect"""
    if aspect not in VALID_ASPECTS:
        raise HTTPException(status_code=400, detail=f"Invalid aspect. Must be one of: {VALID_ASPECTS}")
    file_path = KEYWORDS_DIR / f"{aspect}.json"
    # Backup original
    backup_path = KEYWORDS_DIR / f"{aspect}.json.bak"
    if file_path.exists():
        import shutil
        shutil.copy2(file_path, backup_path)
    with open(file_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    return {"message": f"Keywords for '{aspect}' updated successfully"}


# ── DLI Anomaly Detection ─────────────────────────────────────────────────────
@router.get("/dli/anomaly-report")
def dli_anomaly_report(db: Session = Depends(get_db), _=Depends(require_admin)):
    """
    Laporan anomali DLI untuk validasi kualitas sistem.
    Mendeteksi:
    - Grade 4 (DLI ≥70%) tapi ada aspek < 70%
    - Grade 1 (DLI <40%) tapi ada aspek > 60%
    - Grade 2/3 (40-70%) tapi ada aspek > 80%
    - Keyword paling jarang muncul di dokumen Grade 4
    """
    from backend.utils.dli_anomaly_detector import detect_anomalies, get_grade

    # Ambil semua prediksi DLI dari DB via SQLAlchemy
    preds = db.query(Prediction).filter(Prediction.dli_score.isnot(None)).all()

    predictions = []
    for p in preds:
        full_data = p.get_full_dli_data() or {}
        predictions.append({
            'id': p.id,
            'file_name': p.file_name or f'ID:{p.id}',
            'dli_score': p.dli_score,
            'dli_category': p.dli_category,
            'mindful_score': p.mindful_score or 0,
            'meaningful_score': p.meaningful_score or 0,
            'joyful_score': p.joyful_score or 0,
            'pedagogis_score': p.pedagogis_score or 0,
            'digital_score': p.digital_score or 0,
            'keywords_found': full_data.get('keywords_found', {}),
            'keyword_statistics': full_data.get('keyword_statistics', {}),
            '_full_data': full_data,
        })

    return detect_anomalies(predictions)
