"""
RPPParser — parses Indonesian lesson plan (RPP) documents to extract planned activities.
Supports PDF and DOCX formats.
"""
from __future__ import annotations
import os
from typing import List

from backend.ai.video_3m.data_models import LessonPlan, PlannedActivity


class RPPParser:
    """
    Parses RPP (Rencana Pelaksanaan Pembelajaran) documents and extracts
    planned teaching activities by scanning for Indonesian keyword patterns.
    """

    ACTIVITY_KEYWORDS = {
        "group_discussion": [
            "diskusi kelompok",
            "kerja kelompok",
            "kolaborasi",
            "berdiskusi",
            "diskusi bersama",
            "kerja sama",
        ],
        "mindful_reflection": [
            "refleksi",
            "mindful",
            "renungan",
            "stop",
            "hening",
            "kesadaran penuh",
            "perenungan",
        ],
        "digital_gamification": [
            "gamifikasi",
            "game",
            "kuis digital",
            "quizizz",
            "kahoot",
            "permainan digital",
            "game edukasi",
        ],
        "lecture": [
            "ceramah",
            "penjelasan guru",
            "presentasi guru",
            "menerangkan",
            "menjelaskan materi",
        ],
        "peer_discussion": [
            "diskusi antar siswa",
            "peer learning",
            "teman sebaya",
            "berpasangan",
            "think pair share",
        ],
        "hands_on_activity": [
            "praktikum",
            "eksperimen",
            "proyek",
            "membuat",
            "berkreasi",
            "aktivitas langsung",
        ],
    }

    ACTIVITY_DESCRIPTIONS = {
        "group_discussion": "Diskusi Kelompok",
        "mindful_reflection": "Refleksi Mindful",
        "digital_gamification": "Gamifikasi Digital",
        "lecture": "Ceramah / Penjelasan Guru",
        "peer_discussion": "Diskusi Antar Siswa",
        "hands_on_activity": "Aktivitas Langsung / Praktikum",
    }

    def parse(self, file_path: str) -> LessonPlan:
        """
        Parse a PDF or DOCX RPP file and extract planned activities.
        Returns LessonPlan with activities list and raw text.
        Raises ValueError for invalid/unreadable files.
        """
        if not os.path.exists(file_path):
            raise ValueError(f"File tidak ditemukan: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            raw_text = self._extract_pdf(file_path)
        elif ext in (".docx", ".doc"):
            raw_text = self._extract_docx(file_path)
        else:
            raise ValueError(
                f"Format file tidak didukung: {ext}. Harap unggah file PDF atau DOCX."
            )

        if not raw_text or not raw_text.strip():
            raise ValueError("Dokumen RPP tidak dapat dibaca atau kosong.")

        activities = self._extract_activities(raw_text)

        return LessonPlan(activities=activities, raw_text=raw_text)

    # ------------------------------------------------------------------
    # Private: Text Extraction
    # ------------------------------------------------------------------

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pdfplumber."""
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Gagal membaca file PDF: {e}")

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            return "\n".join(paragraphs)
        except Exception as e:
            raise ValueError(f"Gagal membaca file DOCX: {e}")

    # ------------------------------------------------------------------
    # Private: Activity Extraction
    # ------------------------------------------------------------------

    def _extract_activities(self, text: str) -> List[PlannedActivity]:
        """
        Scan text for activity keywords and return list of PlannedActivity.
        Returns empty list (not an error) when no keywords are found.
        """
        text_lower = text.lower()
        activities: List[PlannedActivity] = []
        found_types = set()

        for activity_type, keywords in self.ACTIVITY_KEYWORDS.items():
            for keyword in keywords:
                if keyword in text_lower and activity_type not in found_types:
                    found_types.add(activity_type)
                    activities.append(PlannedActivity(
                        activity_type=activity_type,
                        description=self.ACTIVITY_DESCRIPTIONS.get(
                            activity_type, activity_type
                        ),
                    ))
                    break  # one match per activity type is enough

        return activities
